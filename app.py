from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import os
import re
import json
import tempfile
from typing import Dict, List, Set
from collections import defaultdict
from flask_cors import CORS
import PyPDF2
from docx import Document
import numpy as np
from sentence_transformers import SentenceTransformer
def cosine_similarity(embeddings1, embeddings2=None):
    """Compute cosine similarity between two sets of embeddings.
    
    Args:
        embeddings1: numpy array of shape (n_samples1, n_features)
        embeddings2: numpy array of shape (n_samples2, n_features). If None, computes similarity with itself.
    
    Returns:
        numpy array of shape (n_samples1, n_samples2) with cosine similarities
    """
    if embeddings2 is None:
        embeddings2 = embeddings1
    
    # Normalize the embeddings
    norm1 = np.linalg.norm(embeddings1, axis=1, keepdims=True)
    norm2 = np.linalg.norm(embeddings2, axis=1, keepdims=True)
    
    # Avoid division by zero
    norm1 = np.where(norm1 == 0, 1e-10, norm1)
    norm2 = np.where(norm2 == 0, 1e-10, norm2)
    
    # Normalize
    embeddings1 = embeddings1 / norm1
    embeddings2 = embeddings2 / norm2
    
    # Compute dot product
    return np.dot(embeddings1, embeddings2.T)
import tempfile
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk import pos_tag, ne_chunk
from nltk.tree import Tree
import string
from word2number import w2n
from datetime import datetime
from dateutil import parser
from transformers import pipeline
import torch
from concurrent.futures import ThreadPoolExecutor
import numpy as np
from collections import defaultdict

# Initialize NLTK and sentence transformer
print("\n=== INITIALIZING COMPONENTS ===")
try:
    # Download required NLTK data
    print("\nChecking NLTK data...")
    nltk_data = {
        'punkt': 'tokenizers/punkt',
        'stopwords': 'corpora/stopwords',
        'wordnet': 'corpora/wordnet',
        'averaged_perceptron_tagger': 'taggers/averaged_perceptron_tagger',
        'maxent_ne_chunker': 'chunkers/maxent_ne_chunker',
        'words': 'corpora/words'
    }
    
    nltk_available = True
    for name, path in nltk_data.items():
        try:
            nltk.data.find(path)
            print(f"✓ NLTK {name} data found")
        except LookupError:
            print(f"  - Downloading NLTK {name} data...")
            try:
                nltk.download(name, quiet=True)
                print(f"✓ Successfully downloaded NLTK {name} data")
            except Exception as e:
                print(f"✗ Failed to download NLTK {name} data: {str(e)}")
                nltk_available = False
    
    # Initialize NLTK components
    if nltk_available:
        print("\nInitializing NLTK components...")
        try:
            stop_words = set(stopwords.words('english'))
            lemmatizer = WordNetLemmatizer()
            print("✓ NLTK components initialized successfully")
        except Exception as e:
            print(f"✗ Failed to initialize NLTK components: {str(e)}")
            nltk_available = False
            stop_words = set()
            lemmatizer = WordNetLemmatizer()  # Initialize with default anyway
    else:
        print("\nSkipping NLTK initialization due to missing data")
        stop_words = set()
        lemmatizer = WordNetLemmatizer()
    
    # Initialize sentence transformer model
    sentence_model = None
    transformer_available = False
    try:
        print("\nInitializing Sentence Transformer...")
        
        # First try to load the fine-tuned model
        fine_tuned_path = "./fine_tuned_skill_model"
        if os.path.exists(fine_tuned_path) and os.path.isdir(fine_tuned_path):
            print("  - Found fine-tuned model, loading...")
            try:
                sentence_model = SentenceTransformer(fine_tuned_path)
                print("  ✓ Loaded fine-tuned model successfully")
                transformer_available = True
            except Exception as e:
                print(f"  ✗ Failed to load fine-tuned model: {str(e)}")
                print("  - Falling back to base model...")
        
        # If fine-tuned model not available or failed to load, use base model
        if not transformer_available:
            print("  - Loading base model (all-MiniLM-L6-v2)...")
            sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            transformer_available = True
            print("  ✓ Loaded base model successfully")
        
        print("✓ Sentence Transformer initialized successfully")
        print(f"  - Model: {sentence_model._get_name()}")
        print(f"  - Device: {next(sentence_model.parameters()).device}")
        
    except Exception as e:
        print(f"✗ Failed to initialize Sentence Transformer: {str(e)}")
        print("  - Semantic similarity features will be disabled")
        transformer_available = False
    
    # Initialize skill embeddings
    skill_embeddings = {}
    if transformer_available and sentence_model:
        print("\nInitializing skill embeddings...")
        try:
            skill_count = sum(len(skills) for skills in SKILL_DB.values())
            print(f"Encoding {skill_count} skills...")
            
            for i, (category, skills) in enumerate(SKILL_DB.items(), 1):
                print(f"  - Processing category {i}/{len(SKILL_DB)}: {category}")
                for skill in skills:
                    try:
                        skill_embeddings[skill] = sentence_model.encode(skill, convert_to_tensor=True)
                    except Exception as e:
                        print(f"    ✗ Error encoding skill '{skill}': {str(e)}")
            
            print(f"✓ Successfully encoded {len(skill_embeddings)}/{skill_count} skills")
        except Exception as e:
            print(f"✗ Failed to initialize skill embeddings: {str(e)}")
            skill_embeddings = {}
    else:
        print("\nSkipping skill embeddings (sentence transformer not available)")
        
except Exception as e:
    print(f"Error during initialization: {str(e)}")
    print("Some features may be limited")
    stop_words = set()
    lemmatizer = WordNetLemmatizer()
    sentence_model = None
    skill_embeddings = {}

app = Flask(__name__)
CORS(app)

@app.route('/')
def index():
    return render_template('index.html')

def clean_text(text):
    """Remove leading unwanted symbols and spaces."""
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    return text

def extract_email(text):
    """Extract the first valid email address from the text."""
    try:
        # First, try to find email in lines containing email-related keywords
        lines = text.lower().split('\n')
        for i, line in enumerate(lines):
            if any(keyword in line for keyword in ['email', 'e-mail', 'mail', '@']):
                # Search this line and the next few lines
                search_text = '\n'.join(lines[i:i+3])
                email_pattern = r'(?i)([a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,})'
                matches = re.findall(email_pattern, search_text)
                if matches:
                    return matches[0].strip()
        
        # If no email found in email-related sections, search entire text
        email_pattern = r'(?i)([a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,})'
        matches = re.findall(email_pattern, text)
        if matches:
            return matches[0].strip()
        
        print("Warning: No email found in text")
        return None
    except Exception as e:
        print(f"Error in email extraction: {str(e)}")
        return None

def extract_phone(text):
    """Extract the first valid phone number from the text."""
    phone_pattern = r'\+?\d{1,3}[-\s]?\(?\d{2,4}\)?[-\s]?\d{3,4}[-\s]?\d{3,4}'
    matches = re.findall(phone_pattern, text)
    for match in matches:
        phone = clean_text(match)
        phone = re.sub(r'[^\d+]', '', phone)
        if phone.startswith('+'):
            return phone
        elif len(phone) == 10:
            return f"{phone[:3]}-{phone[3:6]}-{phone[6:]}"
        elif len(phone) > 10:
            return phone
    return None

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF with better error handling and text cleaning."""
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        text = ''
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            except Exception as e:
                print(f"Error extracting text from PDF page: {str(e)}")
        
        if not text.strip():
            print("Warning: No text extracted from PDF")
            return ''
        
        # Clean up common PDF extraction issues
        text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
        text = text.replace('\x0c', '\n')  # Replace form feed with newline
        text = '\n'.join(line.strip() for line in text.splitlines())
        return text
    except Exception as e:
        print(f"Error in PDF text extraction: {str(e)}")
        return ''

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return text

def parse_resume_with_pyresparser(file_path):
    try:
        # Skip pyresparser and use our own parser
        return None
    except Exception as e:
        print(f"PyResParser error: {str(e)}")
        return None

@app.route('/api/parse-resume', methods=['POST'])
def handle_resume_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
        
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'Empty file name'}), 400
        
    print(f"Processing file: {file.filename}")
    temp_file = None
    
    try:
        # Save the file temporarily
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
        file.save(temp_file.name)
        temp_file.close()  # Close the file before processing
        print(f"File saved temporarily as: {temp_file.name}")
        
        # Extract text based on file type
        try:
            if file.filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(temp_file.name)
                print(f"PDF text extracted, length: {len(text)}")
            elif file.filename.lower().endswith('.docx'):
                text = extract_text_from_docx(temp_file.name)
                print(f"DOCX text extracted, length: {len(text)}")
            else:
                return jsonify({'error': 'Unsupported file format. Please upload PDF or DOCX'}), 400
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return jsonify({'error': f'Error extracting text from file: {str(e)}'}), 500
        
        if not text.strip():
            return jsonify({'error': 'No text could be extracted from the file'}), 400
        
        # Process the extracted text
        try:
            name = extract_name(text)
            print(f"Extracted name: {name}")
            
            email = extract_email(text)
            print(f"Extracted email: {email}")
            
            phone = extract_phone(text)
            print(f"Extracted phone: {phone}")
            
            print("Extracting skills...")
            skills = extract_skills(text)
            print(f"Extracted skills: {skills}")
            
            result = {
                'name': name,
                'email': email,
                'phone': phone,
                'skills': skills
            }
            
            return jsonify(result)
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"Error processing text: {str(e)}\n{error_trace}")
            return jsonify({'error': f'Error processing resume text: {str(e)}'}), 500
            
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Unexpected error: {str(e)}\n{error_trace}")
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500
        
    finally:
        # Clean up temporary file
        if temp_file and os.path.exists(temp_file.name):
            try:
                os.unlink(temp_file.name)
                print(f"Temporary file deleted: {temp_file.name}")
            except Exception as e:
                print(f"Warning: Could not delete temporary file {temp_file.name}: {str(e)}")

def extract_name(text):
    """Extract full name using multiple approaches."""
    try:
        # Get the first few chunks of text where names are typically found
        first_chunk = '\n'.join(text.split('\n')[:5])  # Look at first 5 lines
        
        # Method 1: Try to find name after common resume header patterns
        name_patterns = [
            r'(?i)name\s*[:-]?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})',
            r'(?i)^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*$',
            r'(?i)^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*[,\n]'
        ]
        
        for pattern in name_patterns:
            matches = re.findall(pattern, first_chunk)
            if matches:
                candidate = matches[0].strip()
                if len(candidate.split()) >= 2:
                    print(f"Found name using pattern matching: {candidate}")
                    return candidate
        
        # Method 2: Use NLTK's NER
        tokens = word_tokenize(first_chunk)
        pos_tags = pos_tag(tokens)
        ne_chunks = ne_chunk(pos_tags)
        
        for chunk in ne_chunks:
            if hasattr(chunk, 'label'):
                # This is a named entity
                name = ' '.join(c[0] for c in chunk)
                name = clean_text(name)
                if len(name.split()) >= 2:
                    print(f"Found name using NER: {name}")
                    return name
        
        # Method 3: Look for name-like patterns in the first few lines
        lines = [line.strip() for line in first_chunk.split('\n')]
        for line in lines:
            # Skip lines that are clearly not names
            if any(skip in line.lower() for skip in ['resume', 'cv', 'curriculum', 'vitae', '@', 'email', 'phone', 'address', 'http', 'www']):
                continue
            
            # Clean and check the line
            clean_line = ' '.join(word for word in line.split() if word.isalpha())
            words = clean_line.split()
            
            if 2 <= len(words) <= 4 and all(word[0].isupper() for word in words):
                print(f"Found name using fallback method: {clean_line}")
                return clean_line
        
        print("Warning: No name found in text")
        return None
    except Exception as e:
        print(f"Error in name extraction: {str(e)}")
        return None

# Dictionary of common abbreviations and their full forms
skill_abbreviations = {
    'ML': 'Machine Learning',
    'AI': 'Artificial Intelligence',
    'DL': 'Deep Learning',
    'NLP': 'Natural Language Processing',
    'CV': 'Computer Vision',
    'JS': 'JavaScript',
    'TS': 'TypeScript',
    'BE': 'Backend',
    'FE': 'Frontend',
    'FS': 'Full Stack',
    'DB': 'Database',
    'UI': 'User Interface',
    'UX': 'User Experience',
    'CI': 'Continuous Integration',
    'CD': 'Continuous Deployment',
    'AWS': 'Amazon Web Services',
    'GCP': 'Google Cloud Platform',
    'K8s': 'Kubernetes',
}

# Extended skill patterns with categories
skill_patterns = {
    'Programming Languages': [
        'Python', 'JavaScript', 'Java', 'C++', 'C#', 'Ruby', 'PHP', 'Swift',
        'TypeScript', 'Go', 'Rust', 'Kotlin', 'Scala', 'R', 'MATLAB'
    ],
    'Web Technologies': [
        'React', 'Angular', 'Vue.js', 'Node.js', 'Django', 'Flask', 'FastAPI',
        'HTML', 'CSS', 'SASS', 'LESS', 'Bootstrap', 'Tailwind', 'jQuery',
        'REST API', 'GraphQL', 'WebSocket', 'Express.js', 'Spring Boot'
    ],
    'Databases': [
        'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Elasticsearch',
        'Oracle', 'SQLite', 'Cassandra', 'DynamoDB', 'Firebase'
    ],
    'Cloud & DevOps': [
        'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'Git',
        'CI/CD', 'Terraform', 'Ansible', 'Linux', 'Nginx', 'Apache'
    ],
    'AI & Data Science': [
        'Machine Learning', 'Deep Learning', 'Natural Language Processing',
        'Computer Vision', 'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas',
        'NumPy', 'Data Analysis', 'Big Data', 'Hadoop', 'Spark'
    ],
    'Methodologies': [
        'Agile', 'Scrum', 'Kanban', 'DevOps', 'TDD', 'BDD', 'CI/CD',
        'Microservices', 'RESTful', 'Design Patterns', 'OOP'
    ]
}

# Flatten skill patterns for easier lookup
all_skills = set()
for category in skill_patterns.values():
    all_skills.update(category)

def normalize_skill(skill):
    """Normalize skill by converting to lowercase and removing special characters"""
    return ''.join(c.lower() for c in skill if c.isalnum())

def get_skill_similarity(skill1, skill2):
    """Get similarity score between two skills using string similarity"""
    try:
        # Convert skills to lowercase and remove special characters
        skill1 = normalize_skill(skill1).lower()
        skill2 = normalize_skill(skill2).lower()
        
        # Simple string similarity using Jaccard similarity on character n-grams
        def get_ngrams(s, n=3):
            s = f' {s} '  # Add padding
            return [s[i:i+n] for i in range(len(s)-n+1)]
        
        n = 2  # Use bigrams
        ngrams1 = set(get_ngrams(skill1, n))
        ngrams2 = set(get_ngrams(skill2, n))
        
        if not ngrams1 or not ngrams2:
            return 0.0
            
        intersection = len(ngrams1.intersection(ngrams2))
        union = len(ngrams1.union(ngrams2))
        
        # Also consider word overlap
        words1 = set(skill1.split())
        words2 = set(skill2.split())
        word_intersection = len(words1.intersection(words2))
        word_union = len(words1.union(words2))
        
        # Combine n-gram and word overlap scores
        ngram_similarity = intersection / union if union > 0 else 0
        word_similarity = word_intersection / word_union if word_union > 0 else 0
        
        # Weighted average (adjust weights as needed)
        return 0.7 * ngram_similarity + 0.3 * word_similarity
    except Exception as e:
        print(f"Error in get_skill_similarity: {str(e)}")
        return 0.0

# Comprehensive skill database
SKILL_DB = {
    'Programming Languages': {
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift',
        'kotlin', 'go', 'rust', 'scala', 'perl', 'r', 'matlab', 'sql', 'bash', 'powershell'
    },
    'Web Technologies': {
        'html', 'css', 'react', 'angular', 'vue.js', 'node.js', 'express.js', 'django',
        'flask', 'spring', 'asp.net', 'jquery', 'bootstrap', 'tailwind', 'webpack',
        'graphql', 'rest api', 'web services', 'microservices'
    },
    'Databases': {
        'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sql server',
        'sqlite', 'cassandra', 'dynamodb', 'mariadb', 'neo4j', 'firebase', 'nosql'
    },
    'Cloud & DevOps': {
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'terraform',
        'ansible', 'circleci', 'github actions', 'gitlab ci', 'prometheus', 'grafana',
        'devops', 'ci/cd', 'cloud computing'
    },
    'AI & Data Science': {
        'machine learning', 'deep learning', 'neural networks', 'nlp', 'computer vision',
        'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'keras', 'opencv',
        'data analysis', 'data visualization', 'big data', 'hadoop', 'spark'
    },
    'Tools & Methodologies': {
        'git', 'jira', 'agile', 'scrum', 'kanban', 'tdd', 'unit testing', 'ci/cd',
        'rest', 'soap', 'design patterns', 'oop', 'functional programming'
    }
}

# NER label mapping
NER_CATEGORY_MAP = {
    'PRODUCT': {'Programming Languages', 'Web Technologies', 'Databases'},
    'ORG': {'Cloud & DevOps', 'Tools & Methodologies'},
    'GPE': {'Cloud & DevOps'},
    'WORK_OF_ART': {'Web Technologies', 'Programming Languages'}
}

def extract_skills_rule_based(text: str) -> Dict[str, Set[str]]:
    """Extract skills using exact keyword matching and common variations."""
    skills = defaultdict(set)
    text_lower = text.lower()
    
    try:
        # First check abbreviations and expand them
        for abbr, full_form in skill_abbreviations.items():
            if abbr.lower() in text_lower:
                # Find appropriate category for the skill
                for category, category_skills in SKILL_DB.items():
                    if full_form in category_skills:
                        skills[category].add(full_form)
                        break
        
        # Then check for exact matches from skill database
        for category, category_skills in SKILL_DB.items():
            for skill in category_skills:
                skill_lower = skill.lower()
                # Check for exact match or match with common separators
                if any(pattern in text_lower for pattern in [
                    skill_lower,
                    skill_lower.replace(' ', '-'),
                    skill_lower.replace(' ', '_'),
                    skill_lower.replace(' ', '')
                ]):
                    skills[category].add(skill)
        
        return skills
    except Exception as e:
        print(f"Error in rule-based extraction: {str(e)}")
        return defaultdict(set)

def extract_skills_ner(text: str):
    """Extract skills using NLTK's NER and POS tagging."""
    print("\n=== NER-BASED EXTRACTION STARTED ===")
    try:
        # Ensure NLTK data is available
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('taggers/averaged_perceptron_tagger')
            nltk.data.find('chunkers/maxent_ne_chunker')
        except LookupError as e:
            print("Downloading required NLTK data...")
            nltk.download('punkt', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
            nltk.download('maxent_ne_chunker', quiet=True)
            nltk.download('words', quiet=True)
        
        # Tokenize and tag parts of speech
        tokens = word_tokenize(text)
        pos_tags = pos_tag(tokens)
        
        # Use NLTK's named entity chunker
        ne_chunks = ne_chunk(pos_tags)
        skills = defaultdict(set)
        
        # Extract named entities and noun phrases
        for chunk in ne_chunks:
            if hasattr(chunk, 'label'):
                # This is a named entity
                skill = ' '.join(c[0] for c in chunk)
                skill = clean_text(skill)
                if 1 <= len(skill.split()) <= 3:
                    # Categorize the skill
                    skill_lower = skill.lower()
                    categorized = False
                    for category, category_skills in SKILL_DB.items():
                        if any(s.lower() in skill_lower or skill_lower in s.lower() for s in category_skills):
                            skills[category].add(skill)
                            categorized = True
                    
                    # If no category found, add to a general skills category
                    if not categorized:
                        skills['Other Skills'].add(skill)
        
        print(f"Extracted {sum(len(s) for s in skills.values())} skills using NER")
        return skills
    except Exception as e:
        print(f"Error in NER extraction: {str(e)}")
        return defaultdict(set)

def find_similar_skills(text: str, threshold: float = 0.5) -> Dict[str, Set[str]]:
    """Detect skills using sentence embeddings for semantic similarity."""
    print("\n=== SEMANTIC SIMILARITY STARTED ===")
    global sentence_model, skill_embeddings
    
    # If sentence model is not available, try to initialize it
    if sentence_model is None:
        try:
            print("Initializing sentence transformer model...")
            sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            print("Sentence transformer model loaded successfully")
            
            # Initialize skill embeddings if needed
            if not skill_embeddings and sentence_model:
                print("Initializing skill embeddings...")
                skill_embeddings = {}
                for category, skills in SKILL_DB.items():
                    for skill in skills:
                        try:
                            skill_embeddings[skill] = sentence_model.encode(skill, convert_to_tensor=True)
                        except Exception as e:
                            print(f"Error encoding skill '{skill}': {str(e)}")
        except Exception as e:
            print(f"Could not initialize sentence transformer: {str(e)}")
            print("Semantic similarity features will be disabled")
            return {}
    
    if not sentence_model:
        return {}
        
    skills = defaultdict(set)
    try:
        # Ensure NLTK data is available
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt', quiet=True)
        
        # Split text into sentences using NLTK
        sentences = sent_tokenize(text)
        cleaned_sentences = []
        
        # Clean and validate sentences
        for sent in sentences:
            # Convert to lowercase for better matching
            sent = sent.lower()
            # Only include sentences that might contain skills (containing technical terms or skill-related keywords)
            if any(keyword in sent for keyword in ['experience', 'skill', 'technology', 'proficient', 'knowledge', 'expert', 'familiar', 'worked', 'using', 'development']):
                cleaned_sentences.append(sent)
        
        if not cleaned_sentences:
            return {}
        
        # Get embeddings for relevant sentences
        sentence_embeddings = sentence_model.encode(cleaned_sentences)
        
        # Process each category
        for category, category_skills in SKILL_DB.items():
            if not category_skills:
                continue
                
            # Compare each sentence with skills in this category
            for sent, sent_embedding in zip(cleaned_sentences, sentence_embeddings):
                try:
                    # Get embeddings for skills in this category if not already done
                    if not skill_embeddings:
                        skill_embeddings = {}
                        for skill in category_skills:
                            try:
                                skill_embeddings[skill] = sentence_model.encode(skill, convert_to_tensor=True)
                            except Exception as e:
                                print(f"Error encoding skill '{skill}': {str(e)}")
                    
                    # Get similarities for skills that have embeddings
                    valid_skills = []
                    valid_skill_embeddings = []
                    for skill in category_skills:
                        if skill in skill_embeddings:
                            valid_skills.append(skill)
                            valid_skill_embeddings.append(skill_embeddings[skill])
                    
                    if not valid_skills:
                        continue
                        
                    # Convert embeddings to numpy arrays if they're PyTorch tensors
                    if hasattr(sent_embedding, 'cpu'):
                        sent_embedding = sent_embedding.cpu().numpy()
                    
                    # Stack valid skill embeddings
                    if valid_skill_embeddings and hasattr(valid_skill_embeddings[0], 'cpu'):
                        valid_skill_embeddings = [e.cpu().numpy() for e in valid_skill_embeddings]
                    
                    # Calculate similarities
                    similarities = cosine_similarity(
                        np.array([sent_embedding]),
                        np.array(valid_skill_embeddings)
                    )[0]
                    
                    # Add skills that meet the threshold and are actually mentioned
                    for skill, sim in zip(category_skills, similarities):
                        if sim > threshold and skill.lower() in sent:
                            skills[category].add(skill)
                except Exception as e:
                    print(f"Error comparing skills for category {category}: {str(e)}")
                    continue
    except Exception as e:
        print(f"Error in semantic similarity: {str(e)}")
    
    return skills

def extract_skills(text: str) -> Dict[str, Set[str]]:
    """Hybrid skill extraction combining multiple approaches."""
    print("\n" + "="*50)
    print("STARTING SKILL EXTRACTION")
    print("="*50)
    try:
        print("Starting skill extraction...")
        
        # Step 1: Rule-based extraction (always available)
        print("Performing rule-based extraction...")
        rule_based_skills = extract_skills_rule_based(text)
        
        # Step 2: NER-based extraction (using NLTK)
        print("Performing NER-based extraction...")
        ner_skills = {}
        try:
            ner_skills = extract_skills_ner(text)
        except Exception as e:
            print(f"Error in NER extraction:\n{str(e)}")
            # If NER fails, try to reinitialize NLTK data
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('averaged_perceptron_tagger', quiet=True)
                nltk.download('maxent_ne_chunker', quiet=True)
                nltk.download('words', quiet=True)
                ner_skills = extract_skills_ner(text)  # Try again
            except Exception as e2:
                print(f"Retry failed for NER extraction: {str(e2)}")
        
        # Step 3: Semantic similarity (if model is available)
        similar_skills = {}
        try:
            print("Performing semantic similarity analysis...")
            similar_skills = find_similar_skills(text) or {}
        except Exception as e:
            print(f"Error in semantic similarity: {str(e)}")
        
        # Merge results from all methods
        final_skills = defaultdict(set)
        
        # Always include rule-based skills as the baseline
        for category, skills in rule_based_skills.items():
            if skills:  # Only add non-empty categories
                final_skills[category].update(skills)
        
        # Add NER skills
        for category, skills in ner_skills.items():
            if skills:  # Only add non-empty categories
                final_skills[category].update(skills)
            
        # Add similar skills if available
        for category, skills in (similar_skills or {}).items():
            if skills:  # Only add non-empty categories
                final_skills[category].update(skills)
        
        # Post-process: Remove duplicates and sort
        processed_skills = {}
        for category, skill_set in final_skills.items():
            if not skill_set:  # Skip empty categories
                continue
                
            # Remove skills that are substrings of others
            skills_list = sorted(skill_set, key=len, reverse=True)
            filtered_skills = []
            for skill in skills_list:
                if not any(skill in other and skill != other for other in filtered_skills):
                    filtered_skills.append(skill)
            
            if filtered_skills:  # Only include categories with skills
                processed_skills[category] = sorted(filtered_skills)
        
        print("Skill extraction completed")
        if not processed_skills:
            print("Warning: No skills were extracted")
            
        return processed_skills
        
    except Exception as e:
        print(f"Error in hybrid skill extraction: {str(e)}")
        return {}

def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

@app.route('/api/parse-resume', methods=['POST'])
def parse_resume():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    try:
        # Extract text based on file type
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(file)
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(file)
        else:
            return jsonify({'error': 'Unsupported file format'}), 400

        # Extract information
        result = {
            'name': extract_name(text),
            'email': extract_email(text),
            'phone': extract_phone(text),
            'skills': extract_skills(text)
        }
        
        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
